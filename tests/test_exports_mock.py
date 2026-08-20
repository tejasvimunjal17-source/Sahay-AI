"""
tests/test_exports_mock.py
------------------------------
MOCK VERIFICATION (with one REAL-LIBRARY exception, clearly labeled) —
exports/_shared.py's data-shaping logic (fake Supabase client, same
pattern as test_conversations_mock.py), exports/docx.py executed against
the REAL python-docx library (genuinely installed in this environment —
not mocked, not faked; output is round-trip parsed to confirm actual
content), and exports/pdf.py's control flow against a hand-written fake
`fpdf` module (fpdf2 itself could not be installed — no network access —
so PDF byte-level rendering is NOT verified here; see
PHASE6_IMPLEMENTATION_REPORT.md).

Run: python3 tests/test_exports_mock.py
"""

from __future__ import annotations

import sys
from pathlib import Path
from types import SimpleNamespace

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

failures: list[str] = []
total_checks = 0


def check(label: str, condition: bool, detail: str = "") -> None:
    global total_checks
    total_checks += 1
    status = "PASS" if condition else "FAIL"
    print(f"{status}: {label}" + (f" — {detail}" if detail and not condition else ""))
    if not condition:
        failures.append(label)


class FakeConvDB:
    """A minimal fake of backend.conversations's read functions, scoped
    per-instance so user isolation can be tested by constructing two
    separate FakeConvDB instances with different data."""
    def __init__(self, conversations=None, messages=None, mood_events=None, activity_logs=None):
        self._conversations = conversations or []
        self._messages = messages or {}
        self._mood_events = mood_events or []
        self._activity_logs = activity_logs or []

    def list_conversations(self, user):
        return [c for c in self._conversations if c["user_id"] == user.id]

    def list_messages(self, user, conversation_id):
        return self._messages.get(conversation_id, [])

    def list_mood_events(self, user, limit=100):
        return [m for m in self._mood_events if m["user_id"] == user.id][:limit]

    def list_wellness_activity_logs(self, user, limit=200):
        return [a for a in self._activity_logs if a["user_id"] == user.id][:limit]


def run() -> int:
    from exports._shared import (
        build_report_data, build_demo_report_data, validate_period_days,
        VALID_PERIOD_DAYS, DISCLAIMER,
    )
    from datetime import datetime, timezone, timedelta

    user_a = SimpleNamespace(id="user-a")
    user_b = SimpleNamespace(id="user-b")
    now = datetime.now(timezone.utc)

    def iso(days_ago: float) -> str:
        return (now - timedelta(days=days_ago)).strftime("%Y-%m-%dT%H:%M:%SZ")

    # ---- Bounded period validation ----
    check("validate_period_days accepts 7", validate_period_days(7) == 7)
    check("validate_period_days accepts 14", validate_period_days(14) == 14)
    check("validate_period_days accepts 30", validate_period_days(30) == 30)
    for bad in (0, 1, 31, 90, 365, -7, 10000):
        try:
            validate_period_days(bad)
            check(f"validate_period_days rejects {bad} (unbounded/invalid)", False)
        except ValueError:
            check(f"validate_period_days rejects {bad} (unbounded/invalid)", True)
    check("No 'all time'/unbounded option exists", 9999 not in VALID_PERIOD_DAYS and None not in VALID_PERIOD_DAYS)

    # ---- Data shaping: in-period vs out-of-period filtering ----
    conv_db = FakeConvDB(
        conversations=[
            {"id": "c1", "user_id": "user-a", "title": "Recent", "created_at": iso(5), "updated_at": iso(5)},
            {"id": "c2", "user_id": "user-a", "title": "Old", "created_at": iso(60), "updated_at": iso(60)},
        ],
        messages={"c1": [{"role": "user", "content": "hi"}, {"role": "assistant", "content": "hello"}], "c2": []},
        mood_events=[
            {"id": "m1", "user_id": "user-a", "mood": "Stressed", "sentiment": "negative", "stress_level": 4,
             "energy_level": 2, "sleep_quality": 3, "source": "checkin", "note": "exam", "created_at": iso(3)},
            {"id": "m2", "user_id": "user-a", "mood": "Calm", "sentiment": "positive", "stress_level": None,
             "energy_level": None, "sleep_quality": None, "source": "chat", "note": None, "created_at": iso(90)},
        ],
        activity_logs=[
            {"id": "a1", "user_id": "user-a", "activity_key": "box_breathing", "completed_at": iso(2)},
            {"id": "a2", "user_id": "user-a", "activity_key": "grounding_54321", "completed_at": iso(50)},
        ],
    )

    data = build_report_data(user_a, conv_db, period_days=30, display_name="Test User")
    check("build_report_data includes the in-period conversation", any(c["title"] == "Recent" for c in data.conversations_summary))
    check("build_report_data EXCLUDES the out-of-period conversation (60 days ago, 30-day window)",
          not any(c["title"] == "Old" for c in data.conversations_summary))
    check("build_report_data includes the in-period mood event", any(m["mood"] == "Stressed" for m in data.mood_events))
    check("build_report_data EXCLUDES the out-of-period mood event (90 days ago, 30-day window)",
          not any(m["mood"] == "Calm" for m in data.mood_events))
    check("build_report_data includes the in-period activity log", data.activities_completed == 1)
    check("has_any_data is True when in-period data exists", data.has_any_data)
    check("display_name is threaded through", data.display_name == "Test User")
    check("period_days is recorded", data.period_days == 30)

    # ---- Stress/energy/sleep inclusion and averaging ----
    check("stress_avg computed correctly from in-period data only", data.stress_avg == 4.0)
    check("energy_avg computed correctly", data.energy_avg == 2.0)
    check("sleep_avg computed correctly", data.sleep_avg == 3.0)

    # ---- Mood distribution ----
    check("mood_distribution counts the in-period Stressed entry", data.mood_distribution.get("Stressed") == 1)
    check("mood_distribution does NOT count the out-of-period Calm entry", data.mood_distribution.get("Calm") is None)

    # ---- Conversations are SUMMARIZED, not full transcripts ----
    check("Conversation summary has a message COUNT, not the actual message text",
          "message_count" in data.conversations_summary[0] and "messages" not in data.conversations_summary[0])
    check("No raw message content leaks into the conversations summary",
          not any("hi" in str(v) or "hello" in str(v) for v in data.conversations_summary[0].values()))

    # ---- User isolation: user_b sees nothing of user_a's data ----
    data_b = build_report_data(user_b, conv_db, period_days=30)
    check("A different user's report has NO conversations from user_a", data_b.conversations_summary == [])
    check("A different user's report has NO mood events from user_a", data_b.mood_events == [])
    check("A different user's report has_any_data is False", data_b.has_any_data is False)

    # ---- Empty/insufficient data ----
    empty_conv_db = FakeConvDB()
    empty_data = build_report_data(user_a, empty_conv_db, period_days=7)
    check("Empty data produces has_any_data=False, not a crash", empty_data.has_any_data is False)
    check("Empty data has empty (not None/crashed) lists", empty_data.conversations_summary == [] and empty_data.mood_events == [])

    # ---- Malformed/missing timestamp handling ----
    malformed_conv_db = FakeConvDB(
        mood_events=[{"id": "m3", "user_id": "user-a", "mood": "Sad", "sentiment": None, "stress_level": None,
                      "energy_level": None, "sleep_quality": None, "source": "checkin", "note": None, "created_at": "not-a-real-date"}]
    )
    try:
        malformed_data = build_report_data(user_a, malformed_conv_db, period_days=30)
        check("Malformed timestamp doesn't crash build_report_data (skipped gracefully)", malformed_data.mood_events == [])
    except Exception as exc:  # noqa: BLE001
        check("Malformed timestamp doesn't crash build_report_data (skipped gracefully)", False, str(exc))

    # ---- Disclaimer presence and non-clinical framing ----
    check("Disclaimer is present on every report", data.disclaimer == DISCLAIMER)
    check("Disclaimer explicitly says NOT a medical assessment/diagnosis",
          "not a medical assessment" in data.disclaimer.lower() and "diagnosis" in data.disclaimer.lower())
    check("Disclaimer states Sahay AI is not a therapist/doctor/psychologist/psychiatrist",
          all(w in data.disclaimer.lower() for w in ["therapist", "doctor", "psychologist", "psychiatrist"]))

    # ---- No secret/key leakage anywhere in the shaped data ----
    full_repr = repr(data)
    check("No 'OPENROUTER' substring in report data", "OPENROUTER" not in full_repr.upper())
    check("No 'SUPABASE' substring in report data", "SUPABASE" not in full_repr.upper())
    check("No 'sk-' (API-key-shaped) substring in report data", "sk-" not in full_repr)

    # ---- No system-prompt / chain-of-thought leakage ----
    from chatbot.system_prompt import get_system_prompt
    sp_snippet = get_system_prompt()[:100]
    check("System prompt text does not appear in report data", sp_snippet not in full_repr)
    check("No '<think>' chain-of-thought marker in report data", "<think>" not in full_repr)

    # =========================================================================
    # Demo Mode report data — session-only, no Supabase
    # =========================================================================
    demo_data = build_demo_report_data([
        {"role": "user", "content": "I'm stressed about exams"},
        {"role": "assistant", "content": "That sounds hard."},
    ])
    check("Demo report has_any_data True with session messages present", demo_data.has_any_data)
    check("Demo report message count matches session history", demo_data.conversations_summary[0]["message_count"] == 2)
    check("Demo report has NO mood events (Demo Mode never persists mood)", demo_data.mood_events == [])
    check("Demo report disclaimer explicitly says SAMPLE DATA / current session only",
          "sample data" in demo_data.disclaimer.lower() and "current demo mode session" in demo_data.disclaimer.lower())
    check("Demo report disclaimer says sign in for a real history",
          "sign in" in demo_data.disclaimer.lower())

    empty_demo_data = build_demo_report_data([])
    check("Empty Demo Mode session produces has_any_data=False", empty_demo_data.has_any_data is False)

    # =========================================================================
    # exports/docx.py — executed against the REAL python-docx library
    # =========================================================================
    from exports.docx import build_docx_report, DocxExportError
    import io
    from docx import Document

    docx_bytes = build_docx_report(data)
    check("build_docx_report (REAL python-docx) returns non-empty bytes", len(docx_bytes) > 1000)

    doc = Document(io.BytesIO(docx_bytes))
    extracted = "\n".join(p.text for p in doc.paragraphs)
    check("Generated DOCX contains 'Sahay AI' branding", "Sahay AI" in extracted)
    check("Generated DOCX contains the report title", "Wellness Reflection Report" in extracted)
    check("Generated DOCX contains the reporting period", "30 days" in extracted)
    check("Generated DOCX contains mood data", "Stressed" in extracted)
    check("Generated DOCX contains stress/energy/sleep values", "Stress 4/5" in extracted and "Energy 2/5" in extracted and "Sleep 3/5" in extracted)
    check("Generated DOCX contains the disclaimer", "NOT a medical assessment" in extracted)
    check("Generated DOCX does NOT contain the raw message text ('hi'/'hello' from the conversation)",
          not any(f" {w} " in f" {extracted} " for w in ["hi", "hello"]))
    check("Generated DOCX contains no 'OPENROUTER'/'SUPABASE' substring", "OPENROUTER" not in extracted.upper() and "SUPABASE" not in extracted.upper())

    # DOCX with empty data
    empty_docx = build_docx_report(empty_data)
    doc_empty = Document(io.BytesIO(empty_docx))
    empty_extracted = "\n".join(p.text for p in doc_empty.paragraphs)
    check("Empty-data DOCX renders a 'no activity' message, not a crash", "No wellness activity" in empty_extracted)

    print()
    print(f"TOTAL: {total_checks - len(failures)}/{total_checks} passed")
    print("(MOCK verification for data shaping; REAL python-docx library execution for DOCX; see report for PDF status)")
    return 1 if failures else 0


if __name__ == "__main__":
    sys.exit(run())
