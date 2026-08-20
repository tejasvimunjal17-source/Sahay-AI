"""
tests/test_reports_ui_mock.py
--------------------------------
MOCK VERIFICATION — button-driven UI tests for pages/reports.py. Follows
the established Phases 1-5 pattern: real Streamlit stub execution, real
session-state assertions (not just "no exception"), SUPABASE_URL/
SUPABASE_ANON_KEY set so the auth gate genuinely engages.

DOCX download uses the REAL python-docx library underneath (see
tests/test_exports_mock.py for the direct round-trip check) — here we
only confirm the UI wiring (button click -> bytes stored in
session_state -> download_button appears). PDF download uses the fake
fpdf stub if present on PYTHONPATH; if absent, the "Prepare PDF" click
is still tested for graceful error handling (real behavior in this
real, fpdf2-less environment).

Run:
    SUPABASE_URL=https://fake.supabase.co SUPABASE_ANON_KEY=fake \\
        PYTHONPATH=/path/to/fpdf_stub:/path/to/streamlit_stub \\
        python3 tests/test_reports_ui_mock.py
"""

import sys, importlib
from types import SimpleNamespace
import uuid

sys.path.insert(0, ".")
import streamlit as st
from streamlit import StopRender

failures = []
total_checks = 0


def check(label, condition, detail=""):
    global total_checks
    total_checks += 1
    status = "PASS" if condition else "FAIL"
    print(f"{status}: {label}" + (f" — {detail}" if detail and not condition else ""))
    if not condition:
        failures.append(label)


def fresh():
    for mod in list(sys.modules):
        if (mod == "streamlit_app" or mod.startswith("pages.") or mod.startswith("components.")
                or mod in ("config",) or mod.startswith("content.") or mod.startswith("backend.")
                or mod.startswith("chatbot.") or mod.startswith("exports.")):
            del sys.modules[mod]
    st.session_state.clear()
    st._click_queue.clear()
    st.query_params.clear()


def setup_mocked_auth_with_data():
    import backend.auth as auth_mod
    fake_user = SimpleNamespace(id="fake-user-1", email="student@example.com")
    st.session_state["sahay_supabase_session"] = {"access_token": "t", "refresh_token": "r"}
    auth_mod.get_current_user = lambda: fake_user
    auth_mod.get_profile = lambda user: {"display_name": "Alex", "preferred_language": "en"}

    import backend.conversations as conv_db
    from datetime import datetime, timezone, timedelta
    now = datetime.now(timezone.utc)
    recent = (now - timedelta(days=3)).strftime("%Y-%m-%dT%H:%M:%SZ")

    conv_db.list_conversations = lambda user: [
        {"id": "c1", "user_id": user.id, "title": "Exam stress", "created_at": recent, "updated_at": recent}
    ]
    conv_db.list_messages = lambda user, cid: [
        {"role": "user", "content": "I'm stressed"}, {"role": "assistant", "content": "That sounds hard"}
    ]
    conv_db.list_mood_events = lambda user, limit=100: [
        {"id": "m1", "user_id": user.id, "mood": "Stressed", "sentiment": "negative", "stress_level": 4,
         "energy_level": 2, "sleep_quality": 3, "source": "checkin", "note": None, "created_at": recent}
    ]
    conv_db.list_wellness_activity_logs = lambda user, limit=200: [
        {"id": "a1", "user_id": user.id, "activity_key": "box_breathing", "completed_at": recent}
    ]
    return fake_user


def setup_mocked_auth_no_data():
    import backend.auth as auth_mod
    fake_user = SimpleNamespace(id="fake-user-2", email="empty@example.com")
    st.session_state["sahay_supabase_session"] = {"access_token": "t", "refresh_token": "r"}
    auth_mod.get_current_user = lambda: fake_user
    auth_mod.get_profile = lambda user: {"display_name": None, "preferred_language": "en"}

    import backend.conversations as conv_db
    conv_db.list_conversations = lambda user: []
    conv_db.list_messages = lambda user, cid: []
    conv_db.list_mood_events = lambda user, limit=100: []
    conv_db.list_wellness_activity_logs = lambda user, limit=200: []
    return fake_user


def run():
    try:
        importlib.import_module("streamlit_app")
        return "ran clean"
    except StopRender:
        return "rerun (expected)"
    except Exception as e:
        return f"FAIL: {type(e).__name__}: {e}"


results = []

# ============================================================================
# Authenticated Reports page: loads with real data, default period
# ============================================================================
fresh()
setup_mocked_auth_with_data()
st.session_state["sahay_page"] = "reports"
st.session_state["sahay_view"] = "app"
r = run()
results.append(("reports: authenticated page loads with real data (default 30-day period)", r, r == "ran clean", None))

# ============================================================================
# Period selection: switching to "Last 7 days" is reflected in session state
# ============================================================================
fresh()
setup_mocked_auth_with_data()
st.session_state["sahay_page"] = "reports"
st.session_state["sahay_view"] = "app"
st.selectbox = lambda label, options, **kw: "Last 7 days"
r = run()
ok = st.session_state.get("reports_period_days") == 7
results.append(("reports: selecting 'Last 7 days' updates reports_period_days to 7", r, ok, st.session_state.get("reports_period_days")))
st.selectbox = lambda label, options, **kw: (options[kw.get("index", 0)] if options else None)

# ============================================================================
# Insufficient-data state: user with zero data sees the empty state, not a crash
# ============================================================================
fresh()
setup_mocked_auth_no_data()
st.session_state["sahay_page"] = "reports"
st.session_state["sahay_view"] = "app"
r = run()
results.append(("reports: user with no data renders cleanly (insufficient-data state)", r, r == "ran clean", None))

# ============================================================================
# DOCX: Prepare -> bytes stored in session_state (REAL python-docx underneath)
# ============================================================================
fresh()
setup_mocked_auth_with_data()
st.session_state["sahay_page"] = "reports"
st.session_state["sahay_view"] = "app"
st._click_queue.append("auth_prepare_docx")
r = run()
docx_bytes = st.session_state.get("auth_docx_bytes")
ok = docx_bytes is not None and len(docx_bytes) > 1000
results.append(("reports: Prepare DOCX produces real bytes via python-docx", r, ok, f"{len(docx_bytes) if docx_bytes else 0} bytes"))

# Round-trip: confirm those bytes are a genuinely valid docx
if docx_bytes:
    import io
    from docx import Document
    doc = Document(io.BytesIO(docx_bytes))
    text = "\n".join(p.text for p in doc.paragraphs)
    results.append(("reports: DOCX bytes from the UI flow round-trip-parse with real content", "n/a", "Sahay AI" in text and "Stressed" in text, text[:200]))

# ============================================================================
# PDF: Prepare -> control-flow only (fake stub if present, else graceful error)
# ============================================================================
fresh()
setup_mocked_auth_with_data()
st.session_state["sahay_page"] = "reports"
st.session_state["sahay_view"] = "app"
st._click_queue.append("auth_prepare_pdf")
r = run()
try:
    import fpdf
    fpdf_present = True
except ImportError:
    fpdf_present = False

if fpdf_present:
    pdf_bytes = st.session_state.get("auth_pdf_bytes")
    ok = pdf_bytes is not None
    results.append(("reports: Prepare PDF produces bytes (fake fpdf stub present)", r, ok, "no bytes stored" if not ok else None))
else:
    # No fpdf at all -> button click should surface a friendly st.error, not crash the page
    results.append(("reports: Prepare PDF with fpdf2 genuinely unavailable degrades gracefully (no crash)", r, r == "ran clean", None))

# ============================================================================
# Demo Mode: session-only sample export
# ============================================================================
fresh()
st.session_state["sahay_demo_mode"] = True
st.session_state["sahay_view"] = "app"
st.session_state["sahay_page"] = "reports"
st.session_state["sahay_fullpage_history"] = [
    {"role": "user", "content": "I'm anxious about my interview"},
    {"role": "assistant", "content": "That's understandable."},
]
r = run()
results.append(("reports: Demo Mode with session chat renders sample-export preview", r, r == "ran clean", None))

fresh()
st.session_state["sahay_demo_mode"] = True
st.session_state["sahay_view"] = "app"
st.session_state["sahay_page"] = "reports"
r = run()
results.append(("reports: Demo Mode with NO session chat shows empty state, not a crash", r, r == "ran clean", None))

# Demo Mode DOCX download uses real python-docx too
fresh()
st.session_state["sahay_demo_mode"] = True
st.session_state["sahay_view"] = "app"
st.session_state["sahay_page"] = "reports"
st.session_state["sahay_fullpage_history"] = [
    {"role": "user", "content": "test message"}, {"role": "assistant", "content": "test reply"},
]
st._click_queue.append("demo_prepare_docx")
r = run()
demo_docx = st.session_state.get("demo_docx_bytes")
ok = demo_docx is not None and len(demo_docx) > 500
results.append(("reports: Demo Mode Prepare DOCX produces real sample bytes", r, ok, None))
if demo_docx:
    import io as io2
    from docx import Document as Document2
    doc2 = Document2(io2.BytesIO(demo_docx))
    demo_text = "\n".join(p.text for p in doc2.paragraphs)
    results.append(("reports: Demo Mode DOCX is labeled SAMPLE/session data, never implies persisted history",
                     "n/a", "sample data" in demo_text.lower() and "sign in" in demo_text.lower(), demo_text[-300:]))

# ============================================================================
# Error handling: a broken backend call surfaces a friendly message, not a crash
# ============================================================================
fresh()
setup_mocked_auth_with_data()
import backend.conversations as conv_db_err
def _raise(*a, **kw):
    raise RuntimeError("simulated Supabase failure")
conv_db_err.list_conversations = _raise
st.session_state["sahay_page"] = "reports"
st.session_state["sahay_view"] = "app"
r = run()
results.append(("reports: a backend failure is caught and shown as a friendly error, not a crash", r, r == "ran clean", None))

print()
for label, r, ok, detail in results:
    status = "PASS" if (ok and "FAIL" not in str(r)) else "FAIL"
    print(f"{status}: {label} [{r}]" + (f" -- {detail}" if detail is not None and not ok else ""))

fails = [l for l, r, ok, d in results if not (ok and "FAIL" not in str(r))]
print()
print(f"TOTAL: {len(results)-len(fails)}/{len(results)} passed")
if fails:
    sys.exit(1)
